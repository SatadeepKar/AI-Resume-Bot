"""
PDF & DOCX Generator Module
Converts HTML resume content to PDF and optionally generates DOCX files.
"""
import os
import pdfkit
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

from config import settings, BASE_DIR


# ── PDF Generation ───────────────────────────────────────────────────────────

# Path to wkhtmltopdf binary (local installation in project folder)
WKHTMLTOPDF_PATH = os.path.join(str(BASE_DIR), "wkhtmltopdf", "bin", "wkhtmltopdf.exe")

# pdfkit configuration
PDFKIT_CONFIG = pdfkit.configuration(wkhtmltopdf=WKHTMLTOPDF_PATH) if os.path.exists(WKHTMLTOPDF_PATH) else None

# pdfkit options for clean PDF output
PDF_OPTIONS = {
    "page-size": "A4",
    "margin-top": "12mm",
    "margin-right": "12mm",
    "margin-bottom": "12mm",
    "margin-left": "12mm",
    "encoding": "UTF-8",
    "no-outline": None,
    "enable-local-file-access": None,
}


def generate_pdf(html_content: str, output_path: str) -> str:
    """
    Generate a PDF file from HTML content.
    
    Args:
        html_content: Rendered HTML string.
        output_path: Full path for the output PDF file.
    
    Returns:
        Path to the generated PDF file.
    """
    os.makedirs(os.path.dirname(output_path), exist_ok=True)

    try:
        if PDFKIT_CONFIG:
            pdfkit.from_string(html_content, output_path, options=PDF_OPTIONS, configuration=PDFKIT_CONFIG)
        else:
            pdfkit.from_string(html_content, output_path, options=PDF_OPTIONS)
    except Exception as e:
        raise RuntimeError(
            f"PDF generation failed: {e}. "
            "Ensure wkhtmltopdf is installed at wkhtmltopdf/bin/wkhtmltopdf.exe"
        )

    return output_path


# ── DOCX Generation ─────────────────────────────────────────────────────────

def generate_docx(resume_data: dict, output_path: str) -> str:
    """
    Generate a DOCX file from structured resume data.
    
    Args:
        resume_data: Dict with resume fields.
        output_path: Full path for the output DOCX file.
    
    Returns:
        Path to the generated DOCX file.
    """
    os.makedirs(os.path.dirname(output_path), exist_ok=True)

    doc = Document()
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(10)

    # ── Name ──
    name = resume_data.get("name", "Candidate")
    heading = doc.add_heading(name, level=0)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in heading.runs:
        run.font.size = Pt(22)
        run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)

    # ── Contact Info ──
    contact_parts = []
    if resume_data.get("email"):
        contact_parts.append(resume_data["email"])
    if resume_data.get("phone"):
        contact_parts.append(resume_data["phone"])
    if resume_data.get("linkedin"):
        contact_parts.append(resume_data["linkedin"])
    if contact_parts:
        contact_para = doc.add_paragraph(" | ".join(contact_parts))
        contact_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in contact_para.runs:
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    # ── Summary ──
    if resume_data.get("summary"):
        doc.add_heading("Professional Summary", level=1)
        doc.add_paragraph(resume_data["summary"])

    # ── Skills ──
    if resume_data.get("skills"):
        doc.add_heading("Skills", level=1)
        skills_text = " • ".join(resume_data["skills"])
        doc.add_paragraph(skills_text)

    # ── Experience ──
    if resume_data.get("experience"):
        doc.add_heading("Professional Experience", level=1)
        for exp in resume_data["experience"]:
            # Title and company
            title_para = doc.add_paragraph()
            title_run = title_para.add_run(f"{exp.get('title', '')} — {exp.get('company', '')}")
            title_run.bold = True
            title_run.font.size = Pt(11)

            duration = exp.get("duration", "")
            if duration:
                dur_run = title_para.add_run(f"  |  {duration}")
                dur_run.font.size = Pt(9)
                dur_run.font.color.rgb = RGBColor(0x77, 0x77, 0x77)

            # Bullets
            for bullet in exp.get("bullets", []):
                bullet_para = doc.add_paragraph(bullet, style="List Bullet")
                for run in bullet_para.runs:
                    run.font.size = Pt(10)

    # ── Projects ──
    if resume_data.get("projects"):
        doc.add_heading("Projects", level=1)
        for proj in resume_data["projects"]:
            proj_para = doc.add_paragraph()
            proj_run = proj_para.add_run(proj.get("name", ""))
            proj_run.bold = True
            proj_run.font.size = Pt(11)

            techs = proj.get("technologies", [])
            if techs:
                tech_run = proj_para.add_run(f"  [{', '.join(techs)}]")
                tech_run.font.size = Pt(9)
                tech_run.font.color.rgb = RGBColor(0x44, 0x88, 0xCC)

            desc = proj.get("description", "")
            if desc:
                doc.add_paragraph(desc)

            for bullet in proj.get("bullets", []):
                bullet_para = doc.add_paragraph(bullet, style="List Bullet")
                for run in bullet_para.runs:
                    run.font.size = Pt(10)

    # ── Education ──
    if resume_data.get("education"):
        doc.add_heading("Education", level=1)
        for edu in resume_data["education"]:
            edu_para = doc.add_paragraph()
            edu_run = edu_para.add_run(f"{edu.get('degree', '')} — {edu.get('institution', '')}")
            edu_run.bold = True
            year = edu.get("year", "")
            gpa = edu.get("gpa", "")
            if year or gpa:
                details = []
                if year:
                    details.append(year)
                if gpa:
                    details.append(f"GPA: {gpa}")
                edu_para.add_run(f"  |  {' | '.join(details)}")

    # ── Certifications ──
    if resume_data.get("certifications"):
        doc.add_heading("Certifications", level=1)
        for cert in resume_data["certifications"]:
            doc.add_paragraph(cert, style="List Bullet")

    # Save
    doc.save(output_path)
    return output_path
